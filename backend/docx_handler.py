"""
docx_handler.py – Leser og skriver matte.docx.
Bruker win32com når dokumentet er åpent i Word (fil er låst av OneDrive/Word).
Faller tilbake til python-docx når Word ikke kjører.
"""

import re
import threading
import queue as _queue
from pathlib import Path

# ── Dedikert COM STA-tråd ────────────────────────────────────────────────────
# All win32com-kode MÅ kjøre i denne tråden (Word er STA-objekt).

_com_q: _queue.SimpleQueue = _queue.SimpleQueue()

def _com_worker():
    import pythoncom
    import time
    pythoncom.CoInitialize()
    while True:
        try:
            ev, fn, args, out = _com_q.get_nowait()
            try:
                out["v"] = fn(*args)
            except Exception as e:
                out["e"] = e
            ev.set()
        except _queue.Empty:
            # Pump Windows messages så COM cross-apartment kall kan fullføres
            pythoncom.PumpWaitingMessages()
            time.sleep(0.02)

threading.Thread(target=_com_worker, daemon=True, name="COM-STA").start()

def _com(fn, *args, timeout=30):
    """Kjør fn(*args) i COM-tråden og returner resultatet."""
    ev  = threading.Event()
    out = {}
    _com_q.put((ev, fn, args, out))
    ev.wait(timeout)
    if "e" in out:
        raise out["e"]
    return out.get("v")

TRIGGER_RE = re.compile(
    r"[-–—]\s*l[øoö]ss?[\s!.]*$",
    re.IGNORECASE
)
SEP = "─" * 44


def is_trigger(text: str) -> bool:
    s = text.strip()
    if TRIGGER_RE.search(s):
        return True
    last = s[-15:].lower()
    for v in ["-løs", "- løs", "-los", "- los", "–løs", "– løs"]:
        if last.endswith(v):
            return True
    return False


def clean_task_text(text: str) -> str:
    return TRIGGER_RE.sub("", text).strip()


# ── COM-tilgang ──────────────────────────────────────────────────────────────

def _get_word_doc_impl(doc_path):
    """Kjøres i COM STA-tråden."""
    import win32com.client as win32
    word   = win32.GetActiveObject("Word.Application")
    target = Path(doc_path).name.lower()
    for i in range(1, word.Documents.Count + 1):
        doc = word.Documents(i)
        try:
            name = doc.FullName.split("/")[-1].split("\\")[-1].lower()
            if name == target:
                return word, doc
        except Exception:
            pass
    return None, None

def _get_word_doc(doc_path):
    try:
        return _com(_get_word_doc_impl, doc_path)
    except Exception as e:
        print(f"[COM] feil: {e}")
        return None, None


# ── Lese oppgaver ────────────────────────────────────────────────────────────

def read_tasks(doc_path) -> list:
    def _do(doc_path):
        word, com_doc = _get_word_doc_impl(doc_path)
        if com_doc is not None:
            return _read_tasks_com(com_doc)
        return None
    result = None
    try:
        result = _com(_do, doc_path)
    except Exception as e:
        print(f"[COM] read_tasks feil: {e}")
    if result is not None:
        return result
    # Fallback: python-docx
    from docx import Document
    return _read_tasks_docx(Document(str(doc_path)))


def _read_tasks_com(com_doc) -> list:
    tasks = []
    n = com_doc.Paragraphs.Count
    for i in range(1, n + 1):
        text = com_doc.Paragraphs(i).Range.Text.rstrip("\r\n\x07")
        if not is_trigger(text):
            continue
        # Sjekk om allerede løst
        if i < n:
            nxt = com_doc.Paragraphs(i + 1).Range.Text.strip()
            if nxt.startswith("─") or nxt.startswith("═"):
                continue
        task_text = clean_task_text(text)
        if task_text:
            tasks.append({"index": i, "text": task_text})  # 1-indeksert
    return tasks


def _read_tasks_docx(doc) -> list:
    tasks = []
    paras = doc.paragraphs
    for i, para in enumerate(paras):
        text = para.text
        if not is_trigger(text):
            continue
        if i + 1 < len(paras):
            nxt = paras[i + 1].text.strip()
            if nxt.startswith("─") or nxt.startswith("═"):
                continue
        task_text = clean_task_text(text)
        if task_text:
            tasks.append({"index": i, "text": task_text})  # 0-indeksert
    return tasks


def count_tasks(doc_path) -> int:
    return len(read_tasks(doc_path))


# ── Skrive løsninger ─────────────────────────────────────────────────────────

def write_solutions(doc_path, solutions: list):
    def _do(doc_path, solutions):
        word, com_doc = _get_word_doc_impl(doc_path)
        if com_doc is not None:
            _write_solutions_com(word, com_doc, solutions)
            return True
        return False
    try:
        ok = _com(_do, doc_path, solutions)
    except Exception as e:
        print(f"[COM] write feil: {e}")
        ok = False
    if not ok:
        _write_solutions_docx(doc_path, solutions)


def _write_solutions_com(word_app, com_doc, solutions: list):
    """Setter inn løsninger direkte i åpent Word-dokument via COM."""
    for item in sorted(solutions, key=lambda x: x["index"], reverse=True):
        idx      = item["index"]   # 1-indeksert
        solution = item["solution"].strip()

        if idx > com_doc.Paragraphs.Count:
            continue

        para = com_doc.Paragraphs(idx)

        # Bygg hele løsnings-blokken som én tekst (enklest og sikrere)
        lines   = _build_lines(solution)
        block   = "\n" + "\n".join(t for t, _, _ in lines)

        # Sett inn etter paragrafen ved å bruke InsertAfter på paragrafens Range
        para.Range.InsertAfter(block)

        # Formater "Svar:"-linjer i grønt og fet
        for text, bold, green in lines:
            if not (bold or green):
                continue
            find = com_doc.Content.Find
            find.ClearFormatting()
            find.Replacement.ClearFormatting()
            find.Text = text[:40]   # søk på starten av linjen
            find.Replacement.Text = text[:40]
            find.Replacement.Font.Bold  = True
            find.Replacement.Font.Color = 0x006400
            find.Execute(Replace=1)  # wdReplaceOne
    # Word auto-lagrer via OneDrive – ikke kall Save() her


def _build_lines(solution_text: str) -> list:
    """Returnerer [(tekst, fet, grønn), ...]"""
    lines = [(SEP, False, False)]
    for line in solution_text.split("\n"):
        s = line.strip()
        if not s:
            continue
        if s.lower().startswith("svar:"):
            lines.append((s, True, True))
        else:
            lines.append((s, False, False))
    lines.append((SEP, False, False))
    return lines


def _write_solutions_docx(doc_path, solutions: list):
    """Fallback: skriv via python-docx når Word ikke er åpent."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc_path = Path(doc_path)
    doc      = Document(str(doc_path))

    for item in sorted(solutions, key=lambda x: x["index"], reverse=True):
        idx      = item["index"]
        solution = item["solution"].strip()
        if idx >= len(doc.paragraphs):
            continue

        ref_elem = doc.paragraphs[idx]._element
        for text, bold, green in reversed(_build_lines(solution)):
            p   = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(11)
            if bold:
                run.bold = True
            if green:
                run.font.color.rgb = RGBColor(0, 100, 0)
            elif text.startswith("─"):
                run.font.color.rgb = RGBColor(180, 180, 180)
                run.font.size      = Pt(8)
            ref_elem.addnext(p._element)

    doc.save(str(doc_path))


# ── Eksport-hjelper ──────────────────────────────────────────────────────────

def get_all_paragraphs_with_solutions(doc_path) -> list:
    word, com_doc = _get_word_doc(doc_path)
    if com_doc is not None:
        result = []
        for i in range(1, com_doc.Paragraphs.Count + 1):
            text = com_doc.Paragraphs(i).Range.Text.rstrip("\r\n\x07")
            result.append({
                "text":       text,
                "is_sep":     text.strip().startswith("─"),
                "is_trigger": is_trigger(text),
                "bold":       bool(com_doc.Paragraphs(i).Range.Bold),
            })
        return result

    from docx import Document
    doc    = Document(str(doc_path))
    result = []
    for para in doc.paragraphs:
        text = para.text.strip()
        result.append({
            "text":       text,
            "is_sep":     text.startswith("─"),
            "is_trigger": is_trigger(para.text),
            "bold":       any(run.bold for run in para.runs),
        })
    return result
