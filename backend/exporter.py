"""
exporter.py – Lager matte_besvart.docx uten tekniske elementer.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime


def create_submission_copy(source_path: str | Path) -> Path:
    """
    Lager en ren kopi av matte.docx:
    - Beholder oppgaver og løsninger
    - Fjerner separator-linjer og trigger-ord (- løs)
    - Lagrer som matte_besvart.docx (med dato hvis filen finnes)
    """
    source_path = Path(source_path)
    src_doc     = Document(str(source_path))
    new_doc     = Document()

    # Stil
    style = new_doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Tittel
    title = new_doc.add_heading("Matematikk – Besvarelse", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dato = datetime.date.today().strftime("%d.%m.%Y")
    sub  = new_doc.add_paragraph(f"Dato: {dato}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    new_doc.add_paragraph("")

    # Kopier innhold
    import re
    trigger_re = re.compile(r"[-–—]\s*l[øoö]s[s]?[\s!.]*$", re.IGNORECASE)

    for para in src_doc.paragraphs:
        text = para.text.strip()

        # Hopp over separator-linjer
        if text.startswith("─") or text.startswith("═"):
            continue

        # Fjern trigger-ord fra oppgavetekster
        clean_text = trigger_re.sub("", text).strip()
        if not clean_text:
            continue

        new_p  = new_doc.add_paragraph()
        new_run = new_p.add_run(clean_text)

        # Kopiér fet-formatering
        if any(run.bold for run in para.runs):
            new_run.bold = True

        # Grønn farge for svar-linjer
        if clean_text.lower().startswith("svar:"):
            new_run.bold           = True
            new_run.font.color.rgb = RGBColor(0, 100, 0)

    # Bestem output-filnavn
    output_path = source_path.parent / "matte_besvart.docx"
    if output_path.exists():
        ts          = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = source_path.parent / f"matte_besvart_{ts}.docx"

    new_doc.save(str(output_path))
    print(f"[Exporter] Lagret: {output_path}")
    return output_path
