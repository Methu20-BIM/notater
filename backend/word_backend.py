# -*- coding: utf-8 -*-
"""
word_backend.py – Cross-platform Word-automatisering.
  Windows : win32com COM (live redigering uten filskriving)
  Mac     : python-docx + AppleScript (lagre → rediger → reload)

Felles API:
  doc.paragraph_count()
  doc.paragraph_text(i)            # 1-indeksert
  doc.insert_after(i, block)       # block = newline-separert tekst
  doc.set_paragraph_format(i, bold, color, line_spacing, space_after)
  doc.save()
  doc.get_autosave() / set_autosave(bool)
"""

import platform
import subprocess
from pathlib import Path

IS_WIN = platform.system() == "Windows"
IS_MAC = platform.system() == "Darwin"


# ── Mac-hjelper ───────────────────────────────────────────────────────────────
def _applescript(script: str) -> str:
    r = subprocess.run(["osascript", "-e", script],
                       capture_output=True, text=True)
    return r.stdout.strip()


# ── Windows-implementasjon (win32com) ─────────────────────────────────────────
class _WinDoc:
    def __init__(self, word_app, com_doc):
        self._app = word_app
        self._doc = com_doc

    def paragraph_count(self) -> int:
        return self._doc.Paragraphs.Count

    def paragraph_text(self, i: int) -> str:
        return self._doc.Paragraphs(i).Range.Text

    def insert_after(self, i: int, block: str):
        import time
        for _ in range(5):
            try:
                self._doc.Paragraphs(i).Range.InsertAfter(block)
                return
            except Exception:
                time.sleep(2)

    def set_paragraph_format(self, i: int, bold=None, color=None,
                              line_spacing=None, space_after=None):
        para = self._doc.Paragraphs(i)
        if line_spacing is not None:
            para.Format.LineSpacingRule = line_spacing
        if space_after is not None:
            para.Format.SpaceAfter = space_after
        if bold is not None:
            para.Range.Font.Bold = bold
        if color is not None:
            para.Range.Font.Color = color

    def save(self):
        self._doc.Save()

    def get_autosave(self):
        try:
            return self._doc.AutoSaveOn
        except Exception:
            return None

    def set_autosave(self, value):
        try:
            if value is not None:
                self._doc.AutoSaveOn = value
        except Exception:
            pass


def _get_win_doc(target_name: str):
    import pythoncom
    pythoncom.CoInitialize()
    import win32com.client as win32
    word = win32.GetActiveObject("Word.Application")
    for i in range(1, word.Documents.Count + 1):
        d = word.Documents(i)
        name = d.FullName.replace("\\", "/").split("/")[-1].lower()
        if name == target_name.lower():
            return _WinDoc(word, d)
    return None


# ── Mac-implementasjon (python-docx + AppleScript) ───────────────────────────
class _MacDoc:
    def __init__(self, path: Path):
        self.path = path
        # Lagre Word-dokumentet først slik at filen på disk er oppdatert
        _applescript('tell application "Microsoft Word" to save document 1')
        from docx import Document
        self._doc = Document(str(path))

    def paragraph_count(self) -> int:
        return len(self._doc.paragraphs)

    def paragraph_text(self, i: int) -> str:
        # Legg til \r for å matche Windows COM-oppførsel (rstrip i solve_worker fjerner det)
        return self._doc.paragraphs[i - 1].text + "\r"

    def insert_after(self, i: int, block: str):
        """Setter inn newline-separerte linjer som nye paragrafer etter posisjon i."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        lines = block.split("\n")
        ref_elem = self._doc.paragraphs[i - 1]._element

        # Sett inn i omvendt rekkefølge etter samme referanseelement → riktig rekkefølge
        for line in reversed(lines):
            new_p = OxmlElement("w:p")
            if line:
                new_r = OxmlElement("w:r")
                new_t = OxmlElement("w:t")
                new_t.text = line
                if line.startswith(" ") or line.endswith(" "):
                    new_t.set(qn("xml:space"), "preserve")
                new_r.append(new_t)
                new_p.append(new_r)
            ref_elem.addnext(new_p)

    def set_paragraph_format(self, i: int, bold=None, color=None,
                              line_spacing=None, space_after=None):
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_LINE_SPACING

        para = self._doc.paragraphs[i - 1]
        pf = para.paragraph_format

        if line_spacing == 1:   # Windows 1 = 1.5 linjeavstand
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        if space_after is not None:
            pf.space_after = Pt(space_after)

        if bold is not None or color is not None:
            runs = para.runs
            if not runs:
                runs = [para.add_run()]
            for run in runs:
                if bold is not None:
                    run.bold = bold
                if color is not None:
                    # Windows 0xBBGGRR → RGBColor(R, G, B)
                    r = color & 0xFF
                    g = (color >> 8) & 0xFF
                    b = (color >> 16) & 0xFF
                    run.font.color.rgb = RGBColor(r, g, b)

    def save(self):
        self._doc.save(str(self.path))
        # Last inn igjen i Word fra disk
        _applescript('tell application "Microsoft Word" to revert document 1')

    def get_autosave(self):
        return None  # Ikke aktuelt med filbasert redigering

    def set_autosave(self, value):
        pass


def _get_mac_doc(target_name: str):
    from utils import find_matte_docx
    path = find_matte_docx()
    if not path:
        return None
    if Path(path).name.lower() != target_name.lower():
        return None
    return _MacDoc(Path(path))


# ── Offentlig API ─────────────────────────────────────────────────────────────
def get_doc():
    """Returnerer platform-tilpasset doc-wrapper for åpent matte.docx, eller None."""
    from utils import find_matte_docx
    path = find_matte_docx()
    if not path:
        return None
    target = Path(path).name

    if IS_WIN:
        return _get_win_doc(target)
    if IS_MAC:
        return _get_mac_doc(target)
    raise RuntimeError(f"Plattform {platform.system()} ikke støttet")
