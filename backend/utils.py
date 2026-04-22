"""
utils.py – Hjelpefunksjoner: finn matte.docx, åpne i Word, mapper.
"""

import os
import platform
import subprocess
from pathlib import Path


def get_notater_dir() -> Path:
    """Returnerer mappen der Notater er installert."""
    return Path(__file__).parent.parent


def find_matte_docx() -> Path | None:
    """
    Finner matte.docx – søker i:
    1. Samme mappe som Notater-appen
    2. Skrivebordet
    3. Dokumenter-mappen
    4. C:\\WORD
    Oppretter tom matte.docx på skrivebordet hvis ikke funnet.
    """
    filename = "matte.docx"

    search_dirs = [
        get_notater_dir(),
        _get_desktop(),
        Path.home() / "Documents",
        Path.home() / "OneDrive - Osloskolen" / "Skrivebord",
        Path.home() / "OneDrive - Osloskolen" / "Dokumenter",
        Path("C:/WORD"),
        Path.home() / "Skrivebord",
    ]

    for d in search_dirs:
        candidate = d / filename
        if candidate.exists():
            return candidate

    # Opprett tom fil på skrivebordet
    desktop = _get_desktop()
    desktop.mkdir(parents=True, exist_ok=True)
    new_path = desktop / filename
    _create_empty_matte(new_path)
    return new_path


def _get_desktop() -> Path:
    """Returnerer sti til skrivebordet."""
    # Prøv OneDrive-skrivebord (Osloskolen)
    osloskolen = Path.home() / "OneDrive - Osloskolen" / "Skrivebord"
    if osloskolen.exists():
        return osloskolen

    if platform.system() == "Windows":
        return Path.home() / "Desktop"
    elif platform.system() == "Darwin":
        return Path.home() / "Desktop"
    return Path.home()


def _create_empty_matte(path: Path):
    """Oppretter en tom matte.docx med velkomst-tekst."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Matte – Oppgaver og løsninger", level=1)
    doc.add_paragraph("")
    p = doc.add_paragraph(
        "Skriv oppgaven din her og legg til  - løs  på slutten av linjen."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Eksempel:")

    ex = doc.add_paragraph("Deriver f(x) = x^3 + 2x - 5  - løs")
    ex.runs[0].italic = True

    doc.add_paragraph("")
    doc.add_paragraph(
        "Trykk deretter på 'Løs oppgaver' i Notater-panelet i Word, "
        "eller høyreklikk på Notater-ikonet nede til høyre."
    )
    doc.save(str(path))
    print(f"[Utils] Opprettet ny matte.docx: {path}")


def open_in_word(doc_path: Path | str):
    """Åpner dokumentet i Microsoft Word."""
    doc_path = Path(doc_path)
    if not doc_path.exists():
        return

    system = platform.system()
    try:
        if system == "Windows":
            os.startfile(str(doc_path))
        elif system == "Darwin":
            subprocess.run(["open", "-a", "Microsoft Word", str(doc_path)])
        else:
            subprocess.run(["xdg-open", str(doc_path)])
    except Exception as e:
        print(f"[Utils] Klarte ikke åpne Word: {e}")
