# -*- coding: utf-8 -*-
"""
solve_worker_macos.py – macOS-versjon uten win32com/Word.
Leser og skriver matte.docx direkte via python-docx.
"""

import sys
import json
import re
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from docx import Document
from docx.shared import Pt, RGBColor
from utils import find_matte_docx
from solver import solve_task, ensure_ollama_running

MODEL = "deepseek-r1:7b"

TRIG = re.compile(r"[-\u2013\u2014]\s*l[\u00f8o\u00f6]ss?[\s!.]*$", re.IGNORECASE)
SEP  = "─" * 44

SOLUTION_HEADERS = (
    "hva vi skal finne:",
    "matematisk løsning:",
    "geogebra:",
    "geogebra-kontroll:",
    "rimelighetsvurdering:",
    "svar:",
)


def find_tasks(paragraphs: list) -> list:
    tasks = []
    for i, para in enumerate(paragraphs):
        text = para.text
        if not TRIG.search(text):
            continue
        # Allerede løst hvis neste paragraf inneholder en løsningsoverskrift
        if i + 1 < len(paragraphs):
            nxt = paragraphs[i + 1].text.strip().lower()
            if nxt.startswith("─") or any(nxt.startswith(h) for h in SOLUTION_HEADERS):
                continue
        task_text = TRIG.sub("", text).strip()
        if task_text:
            tasks.append({"index": i, "text": task_text})
    return tasks


def build_lines(solution_text: str) -> list:
    lines = [(SEP, False, False)]
    for line in solution_text.strip().split("\n"):
        s = line.strip()
        if not s:
            continue
        if s.lower().startswith("svar:"):
            lines.append((s, True, True))
        elif any(s.lower().startswith(h) for h in SOLUTION_HEADERS):
            lines.append((s, True, False))
        else:
            lines.append((s, False, False))
    lines.append((SEP, False, False))
    return lines


def insert_solution(doc: Document, idx: int, solution_text: str):
    ref_elem = doc.paragraphs[idx]._element
    for text, bold, green in reversed(build_lines(solution_text)):
        p   = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        if bold:
            run.bold = True
        if green:
            run.font.color.rgb = RGBColor(0, 100, 0)
        elif text.startswith("─"):
            run.font.color.rgb = RGBColor(180, 180, 180)
            run.font.size      = Pt(8)
        ref_elem.addnext(p._element)


def main():
    doc_path = find_matte_docx()
    if not doc_path:
        print(json.dumps({"ok": False, "error": "Finner ikke matte.docx"}))
        return

    doc   = Document(str(doc_path))
    tasks = find_tasks(doc.paragraphs)

    if not tasks:
        print(json.dumps({"ok": True, "count": 0}))
        return

    ensure_ollama_running(MODEL)
    solved = []

    for task in sorted(tasks, key=lambda x: x["index"], reverse=True):
        solution = solve_task(task["text"], MODEL)
        if not solution.lower().startswith("feil"):
            insert_solution(doc, task["index"], solution)
            solved.append(task["index"])

    doc.save(str(doc_path))
    print(json.dumps({"ok": True, "count": len(solved)}))


if __name__ == "__main__":
    main()
